"""
Shared pytest fixtures for the file-manager test suite.

Mock strategy
-------------
- LLMClient  : ollama.Client is patched at the module level so every test gets
                deterministic embeddings (fixed 768-dim vector) and streamed
                generation ("Mock answer.").
- ChromaDB   : config.CHROMA_DIR is redirected to a tmp_path so each test starts
                with an empty, isolated vector store.
- config.ROOT_DIR : redirected to tmp_path so FileManager / Writer stay sandboxed.
"""

from __future__ import annotations

import sys
import os
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock
import pytest

# ---------------------------------------------------------------------------
# Make the project root importable regardless of where pytest is invoked from
# ---------------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import config  # noqa: E402  (must come after sys.path fix)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

EMBED_DIM   = 768
FIXED_EMBED = [0.01] * EMBED_DIM   # deterministic embedding for all texts

def _stream(*words):
    """Return an iterator of ollama-style chat chunk dicts."""
    return iter({"message": {"content": w}} for w in words)


# ---------------------------------------------------------------------------
# Core fixtures
# ---------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def isolate_config(tmp_path, monkeypatch):
    """Redirect ROOT_DIR and CHROMA_DIR to isolated tmp dirs for every test."""
    root   = tmp_path / "root"
    chroma = tmp_path / "chroma"
    root.mkdir()
    chroma.mkdir()
    monkeypatch.setattr(config, "ROOT_DIR",   root)
    monkeypatch.setattr(config, "CHROMA_DIR", chroma)
    monkeypatch.setattr(config, "INDEX_MANIFEST", chroma / "manifest.json")
    monkeypatch.setattr(config, "LOG_FILE",   root / "assistant.log")
    yield


@pytest.fixture()
def mock_ollama_client():
    """
    Patch ollama.Client so no real Ollama process is required.

    Returned MagicMock exposes:
      .embed(model, input)           -> EmbedResponse-like with .embeddings [[...]]
      .embeddings(model, prompt)     -> kept for any legacy call paths
      .chat(model, messages, stream) -> iterator of chunk dicts
      .list()                        -> {"models": [...]}

    ollama 0.6.1 changed embed() to return a Pydantic EmbedResponse where
    .embeddings is List[List[float]] (one vector per input string).
    llm_client.embed() now calls client.embed(), not client.embeddings().
    """
    client = MagicMock()

    # ollama 0.6.1+ API: embed() returns an object with .embeddings: List[List[float]]
    embed_response = MagicMock()
    embed_response.embeddings = [FIXED_EMBED]   # outer list = one result per input
    client.embed.return_value = embed_response

    # Legacy fallback (older ollama builds used embeddings())
    client.embeddings.return_value = {"embedding": FIXED_EMBED}

    # chat() must return a fresh iterator on every call so map-reduce tests
    # (which call generate() multiple times) each get a full "Mock answer." stream.
    client.chat.side_effect = lambda **kw: _stream("Mock", " answer", ".")

    client.list.return_value = {
        "models": [{"name": "phi3:mini"}, {"name": "nomic-embed-text"}]
    }
    return client


@pytest.fixture()
def mock_llm(mock_ollama_client):
    """LLMClient instance backed by the mocked Ollama client."""
    from llm_client import LLMClient
    with patch("llm_client.ollama") as mock_mod:
        mock_mod.Client.return_value = mock_ollama_client
        llm = LLMClient()
        llm._get()   # force lazy init so the mock is attached
        yield llm


@pytest.fixture()
def indexer(mock_llm, tmp_path, monkeypatch):
    """Indexer with mocked LLM and ephemeral ChromaDB in tmp dir."""
    import chromadb
    from indexer import Indexer

    # Use an in-memory ChromaDB to avoid any disk state leaking between tests
    ephemeral = chromadb.EphemeralClient()
    col = ephemeral.get_or_create_collection(
        name=config.COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )

    idx = Indexer(llm=mock_llm)
    idx._col = col   # bypass the lazy PersistentClient init
    return idx


@pytest.fixture()
def file_manager(isolate_config):
    """FileManager rooted at the isolated tmp ROOT_DIR."""
    from file_manager import FileManager
    return FileManager(root=config.ROOT_DIR)


@pytest.fixture()
def writer(file_manager):
    """Writer scoped to the isolated root."""
    from writer import Writer
    return Writer(root=config.ROOT_DIR)


@pytest.fixture()
def summarizer(mock_llm):
    from summarizer import Summarizer
    return Summarizer(llm=mock_llm)


# ---------------------------------------------------------------------------
# Synthetic file helpers (used by multiple test modules)
# ---------------------------------------------------------------------------

SHORT_TEXT = (
    "The patient was diagnosed with hypertension in January 2024. "
    "Blood pressure readings averaged 145/92 mmHg over three visits. "
    "Treatment plan includes lifestyle modifications and beta-blockers."
)

LONG_TEXT = (SHORT_TEXT + "\n\n") * 30   # well above CHUNK_SIZE * 4

CODE_TEXT = '''\
def compute_bmi(weight_kg: float, height_m: float) -> float:
    """Return Body Mass Index rounded to 2 decimal places."""
    if height_m <= 0:
        raise ValueError("height_m must be positive")
    return round(weight_kg / height_m ** 2, 2)


def classify_bmi(bmi: float) -> str:
    if bmi < 18.5:
        return "Underweight"
    elif bmi < 25.0:
        return "Normal"
    elif bmi < 30.0:
        return "Overweight"
    return "Obese"
'''

CSV_TEXT = (
    "patient_id,age,systolic_bp,diastolic_bp,diagnosis\n"
    "P001,45,145,92,Hypertension\n"
    "P002,34,118,76,Normal\n"
    "P003,67,160,98,Stage2Hypertension\n"
    "P004,52,130,84,Prehypertension\n"
)

MARKDOWN_TEXT = (
    "# Health Report Q1 2024\n\n"
    "## Summary\nOverall patient outcomes improved by 12% compared to Q4 2023.\n\n"
    "## Key Metrics\n- Average systolic BP: 132 mmHg\n- Readmission rate: 4.2%\n\n"
    "## Recommendations\nIncrease follow-up appointments for Stage 2 patients.\n"
)


@pytest.fixture()
def fixture_dir(tmp_path):
    """Create a directory containing synthetic files of various types."""
    d = tmp_path / "fixtures"
    d.mkdir()

    (d / "short.txt").write_text(SHORT_TEXT,    encoding="utf-8")
    (d / "long.txt").write_text(LONG_TEXT,      encoding="utf-8")
    (d / "code.py").write_text(CODE_TEXT,       encoding="utf-8")
    (d / "data.csv").write_text(CSV_TEXT,       encoding="utf-8")
    (d / "notes.md").write_text(MARKDOWN_TEXT,  encoding="utf-8")

    # Minimal .docx via python-docx (already in project deps)
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Medical Record", level=1)
        doc.add_paragraph(SHORT_TEXT)
        doc.save(str(d / "record.docx"))
    except Exception:
        pass   # skip .docx if python-docx unavailable in test env

    # Unsupported extension -- should be silently skipped by indexer
    (d / "binary.xyz").write_bytes(b"\x00\x01\x02\x03")

    return d
